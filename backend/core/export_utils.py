"""Export utilities — build download blobs. No Streamlit."""

import io
import json
import zipfile
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import jinja2
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, RGBColor

from config import TEMPLATES_DIR
from logging_config import log_warning


def _empty_llm_metadata() -> Dict[str, Any]:
    return {
        "timestamp": None,
        "provider": None,
        "model": None,
        "temperature": None,
        "system_prompt": None,
        "task_prompt": None,
        "tokens_in": None,
        "tokens_out": None,
        "tokens_think": None,
    }


def build_v2_download_schema(v2_json: Dict[str, Any]) -> Dict[str, Any]:
    """Transform an on-disk V2 transcription JSON into a self-contained download document."""
    src_meta = v2_json.get("metadata", {})
    runs = v2_json.get("runs", [])
    harmonizations = v2_json.get("harmonizations", [])
    summaries = v2_json.get("summaries", [])
    ner_results = v2_json.get("ner_results", [])

    metadata = {
        "image_filename": src_meta.get("image_filename"),
        "created_at": src_meta.get("created_at"),
        "citation": src_meta.get("citation"),
        "license": src_meta.get("license"),
        "app_version": src_meta.get("app_version"),
    }

    last_run = runs[-1] if runs else {}
    shared_context = {
        "profile_name": last_run.get("profile_name"),
        "domain_knowledge": last_run.get("domain_prompt"),
        "notes": "",
    }

    transcription_items: List[str] = []
    for run in runs:
        for output in run.get("outputs", []):
            text = output.get("text", "") if isinstance(output, dict) else str(output)
            transcription_items.append(text)

    if runs:
        transcriptions_meta = {
            "timestamp": last_run.get("completed_at"),
            "provider": last_run.get("provider"),
            "model": last_run.get("model"),
            "temperature": last_run.get("temperature"),
            "system_prompt": last_run.get("base_prompt"),
            "task_prompt": last_run.get("domain_prompt"),
            "tokens_in": last_run.get("tokens_in"),
            "tokens_out": last_run.get("tokens_out"),
            "tokens_think": last_run.get("thinking_tokens", 0),
        }
    else:
        transcriptions_meta = _empty_llm_metadata()

    if harmonizations:
        harm = harmonizations[-1]
        harm_tokens = harm.get("tokens_used") or {}
        harmonization_out = {
            "llm_metadata": {
                "timestamp": harm.get("created_at"),
                "provider": harm.get("provider"),
                "model": harm.get("model_used"),
                "temperature": harm.get("temperature"),
                "system_prompt": harm.get("system_prompt"),
                "task_prompt": harm.get("task_prompt"),
                "tokens_in": harm_tokens.get("prompt_tokens"),
                "tokens_out": harm_tokens.get("completion_tokens"),
                "tokens_think": harm_tokens.get("thinking_tokens"),
            },
            "text": harm.get("harmonized_text", ""),
        }
    else:
        harmonization_out = {"llm_metadata": _empty_llm_metadata(), "text": None}

    if summaries:
        summ = summaries[-1]
        summ_tokens = summ.get("tokens_used") or {}
        summary_out = {
            "llm_metadata": {
                "timestamp": summ.get("created_at"),
                "provider": summ.get("provider"),
                "model": summ.get("model_used"),
                "temperature": summ.get("temperature"),
                "system_prompt": summ.get("system_prompt"),
                "task_prompt": summ.get("task_prompt"),
                "tokens_in": summ_tokens.get("prompt_tokens"),
                "tokens_out": summ_tokens.get("completion_tokens"),
                "tokens_think": summ_tokens.get("thinking_tokens"),
            },
            "text": summ.get("summary_text", ""),
        }
    else:
        summary_out = {"llm_metadata": _empty_llm_metadata(), "text": None}

    if ner_results:
        ner = ner_results[-1]
        ner_tokens_total = (ner.get("tokens_usage") or {}).get("total") or {}
        entity_bundle = ner.get("entity_bundle") or {}
        named_entities_out = {
            "llm_metadata": {
                "timestamp": ner.get("created_at"),
                "provider": ner.get("provider"),
                "model": ner.get("model_used"),
                "temperature": ner.get("temperature"),
                "system_prompt": ner.get("system_prompt"),
                "task_prompt": ner.get("task_prompt"),
                "tokens_in": ner_tokens_total.get("input_tokens"),
                "tokens_out": ner_tokens_total.get("output_tokens"),
                "tokens_think": ner_tokens_total.get("thinking_tokens"),
            },
            "items": entity_bundle.get("entities", []),
            "grounding_info": ner.get("grounding_info"),
            "pass2_used": ner.get("pass2_used"),
        }
    else:
        named_entities_out = {
            "llm_metadata": _empty_llm_metadata(),
            "items": [],
            "grounding_info": None,
            "pass2_used": None,
        }

    return {
        "schema_version": "2.0",
        "metadata": metadata,
        "shared_context": shared_context,
        "outputs": {
            "transcriptions": {
                "llm_metadata": transcriptions_meta,
                "items": transcription_items,
            },
            "harmonization": harmonization_out,
            "summary": summary_out,
            "named_entities": named_entities_out,
        },
    }


_jinja_env = jinja2.Environment(
    loader=jinja2.FileSystemLoader(str(TEMPLATES_DIR)),
    trim_blocks=True,
    lstrip_blocks=True,
    finalize=lambda value: "" if value is None else value,
)


def build_txt_export(v2_json: Dict[str, Any], root: str) -> Tuple[bytes, str, str]:
    """Render a human-readable .txt report from the V2 download schema."""
    download_doc = build_v2_download_schema(v2_json)
    template = _jinja_env.get_template("export_txt.j2")
    rendered = template.render(**download_doc)
    return rendered.encode("utf-8"), f"{root}.txt", "text/plain; charset=utf-8"


def build_single_format_export(
    v2_json: Dict[str, Any], root: str, format: str, image_path: Optional[str] = None
) -> Tuple[bytes, str, str]:
    """Build one export file for one root in the given format ('json', 'txt', or 'docx')."""
    if format == "json":
        download_doc = build_v2_download_schema(v2_json)
        data = json.dumps(download_doc, indent=2, ensure_ascii=False).encode("utf-8")
        return data, f"{root}.json", "application/json"
    if format == "txt":
        return build_txt_export(v2_json, root)
    if format == "docx":
        return build_docx_export(v2_json, root, image_path=image_path)
    raise ValueError(f"Unsupported format: {format}")


def build_zip_export(files: List[Tuple[bytes, str]], format: str) -> Tuple[bytes, str, str]:
    """Zip multiple (data, filename) pairs into one transcriptions_<format>.zip."""
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for data, filename in files:
            zf.writestr(filename, data)
    buf.seek(0)
    return buf.read(), f"transcriptions_{format}.zip", "application/zip"


# ---------------------------------------------------------------------------
# DOCX export
# ---------------------------------------------------------------------------

_DOCX_BRAND_COLOR = RGBColor(0xFF, 0x4B, 0x4B)  # frontend --color-primary
_DOCX_CALLOUT_FILL = "FFF0F0"  # frontend --color-primary-light
_DOCX_USABLE_WIDTH_IN = 6.5  # Letter page, 1in margins on each side

_DOCX_ENTITY_TYPE_LABELS = [
    ("person", "People"),
    ("place", "Places"),
    ("organization", "Organizations"),
    ("date", "Dates"),
]


def _docx_setup_document() -> Document:
    """Create a Letter-sized document with 1in margins."""
    doc = Document()
    now = datetime.now(timezone.utc)
    doc.core_properties.created = now
    doc.core_properties.modified = now
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    return doc


def _docx_add_heading(doc: Document, text: str, level: int = 1):
    """Add a heading styled in the app's brand color."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = _DOCX_BRAND_COLOR
    return heading


def _docx_add_kv_table(doc: Document, rows: List[Tuple[str, Any]]) -> None:
    """Render a 2-column 'Label | Value' table. None/empty values render as an em dash."""
    table = doc.add_table(rows=0, cols=2)
    try:
        table.style = "Light Grid Accent 1"
    except Exception:
        table.style = "Table Grid"
    table.autofit = False
    label_width = Inches(2.0)
    value_width = Inches(_DOCX_USABLE_WIDTH_IN - 2.0)
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].width = label_width
        cells[1].width = value_width
        cells[0].paragraphs[0].add_run(label).bold = True
        display = "—" if value is None or value == "" else str(value)
        cells[1].paragraphs[0].add_run(display)


def _docx_llm_metadata_rows(
    meta: Dict[str, Any], extra: Optional[List[Tuple[str, Any]]] = None
) -> List[Tuple[str, Any]]:
    """Build kv-table rows for an llm_metadata dict (excludes system/task prompts)."""
    rows = [
        ("Provider", meta.get("provider")),
        ("Model", meta.get("model")),
        ("Timestamp", meta.get("timestamp")),
        ("Temperature", meta.get("temperature")),
        ("Tokens In", meta.get("tokens_in")),
        ("Tokens Out", meta.get("tokens_out")),
        ("Thinking Tokens", meta.get("tokens_think")),
    ]
    if extra:
        rows.extend(extra)
    return rows


def _docx_add_prompt_block(doc: Document, label: str, text: Optional[str]) -> None:
    """Add a bold-labeled paragraph for a system/task prompt."""
    p = doc.add_paragraph()
    p.add_run(f"{label}: ").bold = True
    p.add_run(text if text else "[none]")


def _docx_set_cell_shading(cell, hex_color: str) -> None:
    """Apply background shading to a table cell (no public python-docx API for this)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _docx_add_callout(doc: Document, text: str) -> None:
    """Render text inside a shaded, bordered single-cell table."""
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(_DOCX_USABLE_WIDTH_IN)
    _docx_set_cell_shading(cell, _DOCX_CALLOUT_FILL)
    cell.paragraphs[0].add_run(text)
    doc.add_paragraph()


def _docx_add_hyperlink(paragraph, url: str, text: str) -> None:
    """Add a clickable hyperlink run to a paragraph (no public python-docx API for this)."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    run.append(r_pr)

    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _docx_add_page_number_footer(doc: Document) -> None:
    """Add a centered 'Page X of Y' field to the document footer."""
    footer = doc.sections[0].footer
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.text = ""

    def add_field(field_code: str) -> None:
        run = paragraph.add_run()
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = field_code
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        run._r.append(begin)
        run._r.append(instr)
        run._r.append(end)

    paragraph.add_run("Page ")
    add_field(" PAGE ")
    paragraph.add_run(" of ")
    add_field(" NUMPAGES ")


def _docx_add_source_image(doc: Document, image_path: Optional[str]) -> None:
    """Embed the source image, scaled to page width. Skips silently if unavailable."""
    if not image_path:
        return
    path = Path(image_path)
    if not path.exists():
        return
    try:
        doc.add_picture(str(path), width=Inches(_DOCX_USABLE_WIDTH_IN))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception:
        return


def _docx_add_entity(doc: Document, entity: Dict[str, Any]) -> None:
    """Render one named entity as a heading, kv table, and free-text fields."""
    canonical = entity.get("canonical") or entity.get("entity_id") or "Unknown"
    _docx_add_heading(doc, f"{canonical} ({entity.get('entity_id', '')})", level=2)

    variants = entity.get("observed_variants") or []
    _docx_add_kv_table(doc, [
        ("Type", entity.get("type")),
        ("Observed Variants", ", ".join(variants) if variants else None),
        ("Confidence", entity.get("confidence")),
        ("Role in Document", entity.get("role_in_document")),
    ])
    for label, key in (
        ("About", "about"),
        ("Disambiguation", "disambiguation"),
        ("Resolution Notes", "resolution_notes"),
    ):
        value = entity.get(key)
        if value:
            p = doc.add_paragraph()
            p.add_run(f"{label}: ").bold = True
            p.add_run(value)


def _docx_add_entity_index_list(doc: Document, entities: List[Dict[str, Any]], predicate) -> None:
    """Render '<canonical> (<entity_id>)' bullets for entities matching predicate, or a fallback."""
    matches = [e for e in entities if predicate(e)]
    if not matches:
        doc.add_paragraph("[None found]")
        return
    for entity in matches:
        doc.add_paragraph(
            f"{entity.get('canonical', '')} ({entity.get('entity_id', '')})", style="List Bullet"
        )


def build_docx_export(
    v2_json: Dict[str, Any], root: str, image_path: Optional[str] = None
) -> Tuple[bytes, str, str]:
    """Render a styled .docx report from the V2 download schema."""
    download_doc = build_v2_download_schema(v2_json)
    metadata = download_doc["metadata"]
    shared_context = download_doc["shared_context"]
    outputs = download_doc["outputs"]

    doc = _docx_setup_document()

    # Header
    _docx_add_heading(doc, "Document Transcription Export", level=0)
    _docx_add_kv_table(doc, [
        ("Image File", metadata.get("image_filename")),
        ("Schema Version", download_doc.get("schema_version")),
        ("Created At", metadata.get("created_at")),
        ("App Version", metadata.get("app_version")),
        ("License", metadata.get("license")),
    ])
    if metadata.get("citation"):
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run("Citation: ").bold = True
        p.add_run(metadata["citation"])

    _docx_add_source_image(doc, image_path)

    # Shared context
    _docx_add_heading(doc, "Shared Context", level=1)
    _docx_add_kv_table(doc, [
        ("Profile Name", shared_context.get("profile_name")),
        ("Domain Knowledge", shared_context.get("domain_knowledge") or "[none]"),
        ("Notes", shared_context.get("notes") or "[none]"),
    ])

    doc.add_page_break()

    # Summary
    _docx_add_heading(doc, "Summary", level=1)
    doc.add_paragraph(outputs["summary"]["text"] or "[No summary generated.]")

    # Harmonized transcription
    _docx_add_heading(doc, "Harmonized Transcription", level=1)
    _docx_add_callout(
        doc, outputs["harmonization"]["text"] or "[No harmonized transcription generated.]"
    )

    doc.add_page_break()

    # Transcription variants
    _docx_add_heading(doc, "Transcription Variants", level=1)
    _docx_add_heading(doc, "Transcription Model", level=2)
    trans_meta = outputs["transcriptions"]["llm_metadata"]
    _docx_add_kv_table(doc, _docx_llm_metadata_rows(trans_meta))
    _docx_add_prompt_block(doc, "System Prompt", trans_meta.get("system_prompt"))
    _docx_add_prompt_block(doc, "Task Prompt", trans_meta.get("task_prompt"))

    items = outputs["transcriptions"]["items"]
    if items:
        for i, item in enumerate(items, start=1):
            _docx_add_heading(doc, f"Variant {i}", level=2)
            doc.add_paragraph(item)
    else:
        doc.add_paragraph("[No transcription variants generated.]")

    doc.add_page_break()

    # Named entities
    _docx_add_heading(doc, "Named Entities", level=1)
    ner = outputs["named_entities"]
    _docx_add_kv_table(doc, _docx_llm_metadata_rows(
        ner["llm_metadata"], extra=[("Pass 2 Used", ner.get("pass2_used"))]
    ))

    entities = ner["items"]
    if entities:
        for entity in entities:
            _docx_add_entity(doc, entity)
    else:
        doc.add_paragraph("[No named entities extracted.]")

    # Entity index by type
    _docx_add_heading(doc, "Entity Index by Type", level=1)
    for type_key, type_label in _DOCX_ENTITY_TYPE_LABELS:
        _docx_add_heading(doc, type_label, level=2)
        _docx_add_entity_index_list(doc, entities, lambda e, k=type_key: e.get("type") == k)
    _docx_add_heading(doc, "Events / Other", level=2)
    _docx_add_entity_index_list(
        doc, entities,
        lambda e: e.get("type") not in ("person", "place", "organization", "date"),
    )

    doc.add_page_break()

    # Grounding information
    _docx_add_heading(doc, "Grounding Information", level=1)
    grounding = ner.get("grounding_info") or {}

    _docx_add_heading(doc, "Web Search Queries", level=2)
    queries = grounding.get("web_search_queries") or []
    if queries:
        for query in queries:
            doc.add_paragraph(query, style="List Bullet")
    else:
        doc.add_paragraph("[None]")

    _docx_add_heading(doc, "Sources", level=2)
    sources = grounding.get("sources") or []
    if sources:
        for source in sources:
            p = doc.add_paragraph()
            p.add_run(f"{source.get('chunk_index', '')}. {source.get('title') or '[untitled]'}").bold = True
            p2 = doc.add_paragraph()
            p2.add_run("URI: ")
            uri = source.get("uri")
            if uri:
                _docx_add_hyperlink(p2, uri, uri)
            else:
                p2.add_run("[no URI]")
    else:
        doc.add_paragraph("[No sources.]")

    doc.add_page_break()

    # Appendix: processing metadata
    _docx_add_heading(doc, "Appendix: Processing Metadata", level=1)
    for sub_label, key in (
        ("Summary Generation", "summary"),
        ("Harmonization", "harmonization"),
        ("Transcription", "transcriptions"),
        ("Named Entity Extraction", "named_entities"),
    ):
        _docx_add_heading(doc, sub_label, level=2)
        meta = outputs[key]["llm_metadata"]
        _docx_add_kv_table(doc, _docx_llm_metadata_rows(meta))
        _docx_add_prompt_block(doc, "System Prompt", meta.get("system_prompt"))
        _docx_add_prompt_block(doc, "Task Prompt", meta.get("task_prompt"))

    _docx_add_page_number_footer(doc)

    buf = io.BytesIO()
    doc.save(buf)
    return (
        buf.getvalue(),
        f"{root}.docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


def collect_audio_files(workspace, tts_cache_dir: Optional[Path] = None) -> List[Dict[str, Any]]:
    """Collect audio files from TTS cache and workspace. workspace is a SessionWorkspace."""
    from audio_pairing import get_audio_extensions

    audio_files = []

    cache_dir = tts_cache_dir or Path("tmp/audio")
    if cache_dir.exists():
        for wav_file in cache_dir.glob("*.wav"):
            metadata_file = wav_file.with_suffix(".json")
            metadata: Dict[str, Any] = {"source": "tts_generated"}
            if metadata_file.exists():
                try:
                    with open(metadata_file, "r", encoding="utf-8") as f:
                        metadata.update(json.load(f))
                except Exception:
                    pass

            original_image_filename = metadata.get("original_image_filename", "")
            if original_image_filename:
                base_name = Path(original_image_filename).stem
                descriptive_name = f"{base_name}.wav"
            else:
                original_text = metadata.get("original_text", "")
                text_preview = " ".join(original_text.split()[:3]) if original_text else "audio"
                text_preview = "".join(c for c in text_preview if c.isalnum() or c.isspace()).strip()
                text_preview = text_preview.replace(" ", "_")[:20]
                voice = metadata.get("voice", "unknown")
                model = metadata.get("model", "tts")
                timestamp = metadata.get("timestamp", "unknown").replace(":", "-")
                descriptive_name = f"generated_{text_preview}_{voice}_{model}_{timestamp}.wav"

            audio_files.append({"path": wav_file, "filename": descriptive_name, "metadata": metadata})

    if workspace and workspace.workspace_path and workspace.workspace_path.exists():
        audio_extensions = get_audio_extensions()
        for audio_file in workspace.workspace_path.iterdir():
            if audio_file.suffix.lower() in audio_extensions:
                metadata = {
                    "source": "uploaded",
                    "original_filename": audio_file.name,
                    "size_bytes": audio_file.stat().st_size,
                }
                audio_files.append({"path": audio_file, "filename": f"uploaded_{audio_file.name}", "metadata": metadata})

    return audio_files


def build_download_blob_with_audio(
    outputs: List[str],
    download_choice: str = "ZIP of .txt files",
    audio_files: Optional[List[Dict[str, Any]]] = None,
) -> Tuple[bytes, str, str]:
    """Build download blob, optionally including pre-collected audio files."""
    if not outputs:
        return b"", "", "application/octet-stream"

    include_audio = bool(audio_files)

    if download_choice == "JSON dictionary" and not include_audio:
        mapping = {str(i + 1): txt for i, txt in enumerate(outputs)}
        data = json.dumps(mapping, ensure_ascii=False, indent=2).encode("utf-8")
        return data, "transcriptions.json", "application/json"

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for i, txt in enumerate(outputs):
            zf.writestr(f"transcriptions/response_{i+1:02d}.txt", txt)

        if include_audio and audio_files:
            for audio_info in audio_files:
                try:
                    with open(audio_info["path"], "rb") as f:
                        audio_data = f.read()
                    zf.writestr(f"audio/{audio_info['filename']}", audio_data)
                    if audio_info.get("metadata"):
                        meta_name = audio_info["filename"].replace(".wav", "_metadata.json")
                        zf.writestr(f"audio/{meta_name}", json.dumps(audio_info["metadata"], indent=2, ensure_ascii=False))
                except Exception as e:
                    log_warning(f"Could not include audio file {audio_info.get('filename')}: {e}")

    buf.seek(0)
    filename = "transcriptions_with_audio.zip" if include_audio else "transcriptions.zip"
    return buf.read(), filename, "application/zip"
